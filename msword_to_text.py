import os
import subprocess
from docx import Document
import win32com.client as win32

def extract_text_from_word(file_path):
    # Check if file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")
    
    # Convert file path to absolute path
    file_path = os.path.abspath(file_path)
    
    # Extract text based on file extension
    if file_path.lower().endswith(".doc"):
        # Convert .doc to .docx using win32com
        word = None
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            new_file_path = file_path + "x"
            doc.SaveAs(new_file_path, FileFormat=16)  # 16 represents the .docx format
            doc.Close()
            # Use python-docx to extract text from the converted .docx file
            docx = Document(new_file_path)
            full_text = ""
            for para in docx.paragraphs:
                full_text += para.text + "\n"
            return full_text
        except Exception as e:
            raise Exception(f"Error converting and extracting text from .doc file: {str(e)}")
        finally:
            if word is not None:
                word.Quit()
    elif file_path.lower().endswith(".docx"):
        # Use python-docx to extract text from .docx file
        doc = Document(file_path)
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
        return full_text
    else:
        raise ValueError("Unsupported file format. Please provide a .doc or .docx file.")